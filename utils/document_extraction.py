"""Functions for extracting text from various document types."""

import os
import html2text
from pypdf import PdfReader
import docx

# Add imports for Azure Document Intelligence
from azure.core.credentials import AzureKeyCredential
from azure.ai.documentintelligence import DocumentIntelligenceClient
from azure.ai.documentintelligence.models import (
    AnalyzeDocumentRequest,
    DocumentContentFormat,
)

# Document Intelligence supported file types
DOCUMENT_INTELLIGENCE_SUPPORTED_FORMATS = [
    ".pdf",
    ".jpeg",
    ".jpg",
    ".png",
    ".bmp",
    ".tiff",
    ".docx",
    ".xlsx",
    ".pptx",
    ".html",
]


def is_document_intelligence_available():
    """Check if Azure Document Intelligence credentials are available."""
    endpoint = os.environ.get("AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT")
    api_key = os.environ.get("AZURE_DOCUMENT_INTELLIGENCE_API_KEY")
    return endpoint is not None and api_key is not None


def extract_using_document_intelligence(file_path):
    """Extract text from a document using Azure Document Intelligence."""
    endpoint = os.environ.get("AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT")
    api_key = os.environ.get("AZURE_DOCUMENT_INTELLIGENCE_API_KEY")

    # Open the file in binary mode
    with open(file_path, "rb") as f:
        document_bytes = f.read()

    document_intelligence_client = DocumentIntelligenceClient(
        endpoint=endpoint, credential=AzureKeyCredential(api_key)
    )

    poller = document_intelligence_client.begin_analyze_document(
        "prebuilt-layout",
        AnalyzeDocumentRequest(bytes_source=document_bytes),
        output_content_format=DocumentContentFormat.MARKDOWN,
    )

    result = poller.result()
    return result.content


def extract_text(file_path, use_document_intelligence=True):
    """Extract text from various document types (PDF, HTML, TXT, DOCX, JSON)"""
    file_extension = os.path.splitext(file_path)[1].lower()

    # Check if Document Intelligence should be used and is available
    if use_document_intelligence and is_document_intelligence_available():
        if file_extension in DOCUMENT_INTELLIGENCE_SUPPORTED_FORMATS:
            try:
                return extract_using_document_intelligence(file_path)
            except Exception as e:
                print(
                    f"Document Intelligence failed: {str(e)}. Falling back to local processing."
                )
                # Continue to local processing methods

    # Local processing methods
    if file_extension == ".pdf":
        return extract_from_pdf(file_path)
    elif file_extension == ".html" or file_extension == ".htm":
        return extract_from_html(file_path)
    elif file_extension == ".txt":
        return extract_from_txt(file_path)
    elif file_extension == ".docx":
        return extract_from_docx(file_path)
    elif file_extension == ".json":
        return extract_from_json(file_path)
    elif file_extension in DOCUMENT_INTELLIGENCE_SUPPORTED_FORMATS:
        return f"This file format ({file_extension}) requires Azure Document Intelligence, which is not available."
    else:
        return f"Unsupported file type: {file_extension}"


def extract_from_pdf(file_path):
    """Extract text from PDF files."""
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n\n"
    return text


def extract_from_html(file_path):
    """Extract text from HTML files."""
    with open(file_path, "r", encoding="utf-8") as file:
        html_content = file.read()

    # Convert HTML to markdown to preserve structure
    h = html2text.HTML2Text()
    h.ignore_links = False
    h.ignore_images = False
    h.ignore_tables = False
    h.body_width = 0  # No wrapping

    markdown_text = h.handle(html_content)
    return markdown_text


def extract_from_txt(file_path):
    """Extract text from TXT files."""
    with open(file_path, "r", encoding="utf-8") as file:
        return file.read()


def extract_from_docx(file_path):
    """Extract text from DOCX files."""
    doc = docx.Document(file_path)
    text = ""

    # Extract headers and preserve their level
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            level = paragraph.style.name.replace("Heading", "")
            try:
                level_num = int(level.strip())
                text += "#" * level_num + " " + paragraph.text + "\n\n"
            except ValueError:
                text += paragraph.text + "\n\n"
        else:
            text += paragraph.text + "\n\n"

    # Handle tables
    for table in doc.tables:
        text += "\n|"
        for cell in table.rows[0].cells:
            text += f" {cell.text} |"
        text += "\n|"
        for _ in table.rows[0].cells:
            text += " --- |"
        text += "\n"

        for row in table.rows[1:]:
            text += "|"
            for cell in row.cells:
                text += f" {cell.text} |"
            text += "\n"
        text += "\n"

    return text


def extract_from_json(file_path):
    """Extract text from JSON files."""
    with open(file_path, "r", encoding="utf-8") as file:
        return file.read()
